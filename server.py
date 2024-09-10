import pandas as pd
import numpy as np
from http.server import BaseHTTPRequestHandler, HTTPServer
import urllib.parse
import os
import pandas as pd
import numpy as np
import contract_grab_json
from docx import Document  # 导入 python-docx 库

# 加载数据
file_path = "/Users/chiachi/Documents/Master/code/data.xls"
df = pd.read_excel(file_path)

class RequestHandler(BaseHTTPRequestHandler):
    def do_GET(self):
        if self.path == '/':
            self.path = 'index.html'
        elif self.path.startswith('/download'):
            self.handle_download()
            return
        
        try:
            with open(self.path, 'rb') as file:
                self.send_response(200)
                self.send_header('Content-type', self.get_content_type(self.path))
                self.end_headers()
                self.wfile.write(file.read())
        except FileNotFoundError:
            self.send_error(404, "File not found")
    
    def handle_download(self):
        # 处理下载请求
        parsed_path = urllib.parse.urlparse(self.path)
        query = urllib.parse.parse_qs(parsed_path.query)
        file_path = query.get('file', [None])[0]

        if file_path and os.path.exists(file_path):
            self.send_response(200)
            self.send_header('Content-type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Disposition', f'attachment; filename="{os.path.basename(file_path)}"')
            self.end_headers()
            with open(file_path, 'rb') as file:
                self.wfile.write(file.read())
        else:
            self.send_response(404)
            self.end_headers()
            self.wfile.write(b"File not found")
    
    def do_POST(self):
        # 读取 POST 请求的数据
        content_length = int(self.headers['Content-Length'])
        post_data = self.rfile.read(content_length).decode('utf-8')
        params = urllib.parse.parse_qs(post_data)
        
        # 检查是否包含 'selected_data'
        selected_data = params.get("selected_data", [""])[0]
        if selected_data:
            # 处理从客户端接收到的选中行数据
            print(f"使用者選擇了: {selected_data}")
            response = f"服務器已接受到您選擇了: {selected_data}"
            
            selected_data_list = selected_data.split(", ")
            contract_data = df[(df["生效合約編號"] == selected_data_list[0]) & (df["市調合約編號"] == selected_data_list[1]) & (df["合約生效日"] == int(selected_data_list[2])) & (df["合約截止日"] == int(selected_data_list[3]))]
            print(type(df["合約生效日"][206]))
            print(type(int(selected_data_list[2])))
            
            if not contract_data.empty:
                print("hi")
                # 部門
                department_conditions = {
                    "全企業": ["通用性/總處督導公司合約", "合約期間內調整合約（排除運輸合約）"],
                    "單一公司": ["NTD 200萬元以上長庚醫院合約", "NTD 200萬元以上單一督導公司合約", "NTD 20萬元以上生醫體系合約"]
                }
                
                results=[]
                def determine_value(value, conditions):
                    for key, condition_values in conditions.items():
                        if value in condition_values:
                            return key
                    return None
                
                # 新舊約
                def determine_contract_type(value):
                    return "新合約" if "*" in value else "舊合約" 
                
                # 材料項次
                def determine_material_type(value):
                    return "單項材料" if value==1 else "多項材料"
                
                # 詢價
                def determine_inquiry_type(value):
                    return "獨家報價" if value==1 else "多家報價"
                
                # 議價
                def determine_negotiation_type(value, contract_type, material_type, inquiry_type):
                    if value=="      ": 
                        return "無上次訂約" #0
                    # elif determine_contract_type(value)=="舊合約" and determine_material_type(value)=="多項材料" and determine_inquiry_type(value)=="獨家報價":
                    elif contract_type == "舊合約" and material_type == "多項材料" and inquiry_type == "獨家報價":
                        return "訂約一家" #2
                    else:
                        return "有上次訂約" #1

                # 金額比較
                def determine_price_comparison_type(value):
                    if value>0.00:
                        return "較前一般採購價高"
                    elif value==0.00:
                        return "同前一般採購價"
                    elif value<0.00:
                        return "較前一般採購價低"
                    else:
                        return "無前一般採購價"


                department = np.vectorize(determine_value)(contract_data["呈核類型名稱"].values[0], department_conditions)
                if department=="全企業":
                    results.append("0")
                elif department=="單一公司":
                    results.append("1")

                contract_type = np.vectorize(determine_contract_type)(contract_data["市調合約編號"].values[0])
                if contract_type=="新合約":
                    results.append("0")
                elif contract_type=="舊合約":
                    results.append("1")
                
                material = np.vectorize(determine_material_type)(contract_data["規格數量"].values[0])
                if material=="單項材料":
                    results.append("0")
                elif material=="多項材料":
                    results.append("1")
                
                inquiry = np.vectorize(determine_inquiry_type)(contract_data["報價廠商數"].values[0])
                if inquiry=="獨家報價":
                    results.append("0")
                elif inquiry=="多家報價":
                    results.append("1")

                negotiation = np.vectorize(determine_negotiation_type)(contract_data["前次訂約廠商簡稱"].values[0], contract_type, material, inquiry)
                if negotiation=="無上次訂約":
                    results.append("0")
                elif negotiation=="有上次訂約":
                    results.append("1")
                else: #訂約一家
                    results.append("2")
                
                price_comparison = np.vectorize(determine_price_comparison_type)(contract_data["議價調幅"].values[0])
                if price_comparison=="較前一般採購價高":
                    results.append("0")
                elif price_comparison=="同前一般採購價":
                    results.append("1")
                elif price_comparison=="較前一般採購價低":
                    results.append("2")
                elif price_comparison=="無前一般採購價":
                    results.append("3")

                print(f"使用部門: {department}")
                print(f"新舊約: {contract_type}")
                print(f"材料項次: {material}")
                print(f"詢價: {inquiry}")
                print(f"議價: {negotiation}")
                print(f"金額比較: {price_comparison}\n")
                #print(results)

                item = 0
                item_mapping = {
                    "000000": 1, "000001": 2, "000002": 3, "000003": 4,
                    "000100": 5, "000101": 6, "000102": 7, "000103": 8,
                    "001000": 9, "001001": 10, "001002": 11, "001003": 12,
                    "001100": 13, "001101": 14, "001102": 15, "001103": 16,
                    "010010": 17, "010011": 18, "010012": 19, "010110": 20,
                    "010111": 21, "010112": 22, "011020": 23, "011021": 24,
                    "011022": 25, "011110": 26, "011111": 27, "011112": 28,
                    "100000": 29, "100001": 30, "100002": 31, "100003": 32,
                    "100100": 33, "100101": 34, "100102": 35, "100103": 36,
                    "101000": 37, "101001": 38, "101002": 39, "101003": 40,
                    "101100": 41, "101101": 42, "101102": 43, "101103": 44,
                    "110010": 45, "110011": 46, "110012": 47, "110110": 48,
                    "110111": 49, "110112": 50, "111020": 51, "111021": 52,
                    "111022": 53, "111110": 54, "111111": 55, "111112": 56
                }
                item = item_mapping.get("".join(results), 0)
                print(f"項次={item}")

                data = contract_grab_json.grab_data_from_excel(file_path, selected_data_list)
                if data:
                    report1 = contract_grab_json.generate_audit_report(item, data)
                    print(f"{report1}\n")
                    #report_html = f"<html><body><h1>Audit Report</h1>{report1}</body></html>"
                    response = report1
                    
                    # 创建一个新的 Word 文档
                    doc = Document()
                    doc.add_heading('審核報告', level=1)
                    doc.add_paragraph(report1)
                    
                    # 保存文档到指定路径
                    word_file_path = "report.docx"
                    doc.save(word_file_path)

                    # 返回文件路径作为下载链接
                    #response2 = f'<a href="/download?file={word_file_path}" download>Download Report</a>'
                   
        else:
            contract_number = params.get("contract_number", [""])[0]
            # 根据生效合约编号筛选数据
            contract_data = df[(df["生效合約編號"] == contract_number) | (df["市調合約編號"] == contract_number)]
            if contract_data.empty:
                response = "找不到相應的生效合約編號"
            else:
                # 只选择指定的列
                selected_columns = ["生效合約編號", "市調合約編號", "合約生效日", "合約截止日"]
                selected_data = contract_data[selected_columns]
                # 将数据转换为 HTML 表格格式
                response = selected_data.to_html(classes='data', header=True, index=False)
                    
        self.send_response(200)
        self.send_header('Content-type', 'text/html')
        self.end_headers()
        self.wfile.write(response.encode('utf-8'))
        #self.wfile.write(response2.encode('utf-8'))
        #return  # 结束处理以避免继续处理其他逻辑

    def get_content_type(self, path):
        ext = os.path.splitext(path)[1]
        if ext == '.html':
            return 'text/html'
        elif ext == '.css':
            return 'text/css'
        elif ext == '.js':
            return 'application/javascript'
        elif ext == '.jpg' or ext == '.jpeg':
            return 'image/jpeg'
        elif ext == '.png':
            return 'image/png'
        else:
            return 'application/octet-stream'

def run(server_class=HTTPServer, handler_class=RequestHandler, port=8000):
    print('123')
    server_address = ('', port)
    try:
        httpd = server_class(server_address, handler_class)
        print(f'啟動服務器：http://localhost:{port}')
        httpd.serve_forever()
    except Exception as e:
        print('456')
        print(f"啟動服務器時出錯: {e}")

if __name__ == "__main__":
    print("???")
    run()
